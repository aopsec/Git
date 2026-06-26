# -*- coding: utf-8 -*-
"""Builds OpenBox_AVFinal.docx from the AVFinal content, mirroring the HTML deliverable."""
import pathlib
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = pathlib.Path(__file__).parent
FIGS = HERE / "figs"

ACCENT   = RGBColor(0x0F, 0x4C, 0x81)
ACCENT2  = RGBColor(0x8A, 0x1F, 0x3D)
ACCENT3  = RGBColor(0x1F, 0x7A, 0x5A)
MUTED    = RGBColor(0x5B, 0x64, 0x70)
INK      = RGBColor(0x16, 0x19, 0x1F)

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for hs, sz, col in [("Heading 1", 20, ACCENT), ("Heading 2", 14, ACCENT), ("Heading 3", 11.5, ACCENT2)]:
    st = doc.styles[hs]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = col
    st.font.bold = True

# page margins
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)

def para(text="", size=None, color=None, bold=False, italic=False, align=None, after=None, before=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    if after is not None: p.paragraph_format.space_after = Pt(after)
    if before is not None: p.paragraph_format.space_before = Pt(before)
    if text:
        r = p.add_run(text)
        if size: r.font.size = Pt(size)
        if color: r.font.color.rgb = color
        r.bold = bold; r.italic = italic
    return p

def rich(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6):
    """parts: list of (text, dict) where dict may have bold/italic/code/color/size."""
    p = doc.add_paragraph(); p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    for text, opt in parts:
        r = p.add_run(text)
        if opt.get("bold"): r.bold = True
        if opt.get("italic"): r.italic = True
        if opt.get("code"):
            r.font.name = "Consolas"; r.font.size = Pt(9.3)
        if opt.get("size"): r.font.size = Pt(opt["size"])
        if opt.get("color"): r.font.color.rgb = opt["color"]
    return p

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        rr = hdr[i].paragraphs[0].add_run(h)
        rr.bold = True; rr.font.size = Pt(9.5)
        shade(hdr[i], "EAEFF6")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(val)
            rr.font.size = Pt(9.3)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

def callout(label, text, kind="note"):
    color = {"note": ACCENT, "warn": ACCENT2, "ok": ACCENT3}[kind]
    fill  = {"note": "F3F7FB", "warn": "FDECEC", "ok": "ECF6F0"}[kind]
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]; shade(c, fill)
    p = c.paragraphs[0]
    r = p.add_run(label + " "); r.bold = True; r.font.color.rgb = color; r.font.size = Pt(10)
    r2 = p.add_run(text); r2.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def figure(name, caption, width_in=6.3):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(FIGS / name), width=Inches(width_in))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    rr = cap.add_run(caption); rr.italic = True; rr.font.size = Pt(8.5); rr.font.color.rgb = MUTED

def h2(text): doc.add_heading(text, level=2)
def h3(text): doc.add_heading(text, level=3)

# =================== CAPA ===================
band = para("", after=4)
band.add_run("AVFinal   ·   UniCEUB · 2026   ·   Introdução à Computação   ·   Infraestrutura & Segurança")
for r in band.runs: r.font.size = Pt(9); r.font.color.rgb = ACCENT; r.bold = True
para("OpenBox v0.2.0", size=30, color=ACCENT, bold=True, after=2)
para("Plataforma de borda com privacidade auditável, hardening e monitoramento integrado",
     size=13.5, color=MUTED, after=10)
rich([
    ("Documentação técnica derivada do repositório ", {}),
    ("OpenBox0.1v", {"code": True}),
    (", unindo a borda de rede ", {}), ("OpenBox", {"bold": True}),
    (" — gateway de privacidade sobre Rockchip RK3229 (placa Shenzhen ", {}),
    ("R29_5G_LP3", {"code": True}),
    (") — e o plano de detecção host-side ", {}), ("ADV7Sec", {"bold": True}),
    (". O trabalho integra história do projeto, análise de risco, procedimentos de instalação, "
     "hardening e validação operacional, em uma arquitetura residencial defensiva orientada por evidência.", {}),
])
figure("fig1_arch.png", "Figura 1 — Resumo visual da arquitetura: defesa em profundidade (rede, DNS, VPN, monitoramento e integridade).")

meta = [
    ("Disciplina", "Introdução à Computação"),
    ("Instituição", "Centro Universitário de Brasília — UniCEUB"),
    ("Autor", "Alcides Olivo Pollazzon Soterio"),
    ("Base técnica", "RK3229 / R29_5G_LP3 · Armbian rk322x-box · WireGuard · nftables · DNSCrypt · Pi-hole"),
    ("Entrega", "Documento acadêmico expandido com figuras, tabelas, runbook e referências"),
    ("Local · Ano", "Brasília – DF · 2026"),
]
mt = doc.add_table(rows=0, cols=2)
for k, v in meta:
    cells = mt.add_row().cells
    rk = cells[0].paragraphs[0].add_run(k); rk.bold = True; rk.font.size = Pt(9.5)
    cells[0].width = Cm(4)
    rv = cells[1].paragraphs[0].add_run(v); rv.font.size = Pt(9.5); rv.font.color.rgb = MUTED

doc.add_page_break()

# =================== RESUMO + SUMÁRIO ===================
h2("Resumo executivo")
rich([
    ("O ", {}), ("OpenBox v0.2.0", {"bold": True}),
    (" é uma plataforma de borda open source retargetada para a placa Shenzhen ", {}),
    ("R29_5G_LP3", {"code": True}),
    (", com SoC Rockchip RK3229, 1 GB LPDDR3 e Ethernet de 100 Mbps. A solução combina Armbian "
     "rk322x-box, WireGuard, nftables, dnscrypt-proxy, Pi-hole, Unbound e Tor — este reservado a "
     "metadados — somados a um conjunto de monitores locais para observabilidade contínua. A evolução "
     "do projeto priorizou correções factuais, verificabilidade e mitigação realista de ameaças. O "
     "objetivo não é prometer anonimato total nem streaming via Tor, mas reduzir risco, vazamento de "
     "DNS, exposição de metadados e impacto de backdoors de cadeia de suprimento.", {}),
])
h2("Sumário")
sumario = [
    "1. Introdução e objetivos", "2. Contexto e justificativa", "3. Linha do tempo e evolução",
    "4. Estudo de caso: TV box e ameaça de fábrica", "5. Arquitetura OpenBox",
    "6. Demanda de hardware, software e rede", "7. Modelo de ameaça e matriz de risco",
    "8. Plano de hardening e mitigação", "9. ADV7Sec e monitoração host-side",
    "10. Runbook, validação e evidências", "11. Evidências de desenvolvimento",
    "12. Conclusão e direcionamento de carreira", "13. Referências técnicas essenciais",
    "14. Checklist final de conformidade",
]
for item in sumario:
    p = doc.add_paragraph(item, style="List Bullet"); p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.size = Pt(10)

doc.add_page_break()

# =================== 1. INTRODUÇÃO ===================
h2("1. Introdução e objetivos")
para("A disciplina Introdução à Computação exige a identificação de um problema real, a especificação "
     "de uma solução técnica e a documentação de suas limitações. No OpenBox, o problema assumido é "
     "particularmente concreto: dispositivos TV box genéricos, baratos e amplamente distribuídos podem "
     "sair de fábrica com firmware comprometido, serviços de telemetria não desejados e interfaces de "
     "administração frágeis. O projeto responde a esse cenário com uma arquitetura enxuta, auditável e "
     "executável em hardware reaproveitado.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
para("A proposta se apoia em dois pilares. O primeiro é o OpenBox, a borda de rede: um gateway que "
     "controla a saída do tráfego, aplica políticas de DNS, mantém VPN com kill switch e segmenta a rede "
     "doméstica. O segundo é o ADV7Sec, um plano de controle orientado à detecção, que observa a estação "
     "de trabalho e agrega eventos de auditd, Falco, Suricata, Zeek e verificadores de integridade. Os "
     "dois elementos se complementam e transformam a casa em um laboratório de defesa realista.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
h3("Objetivos específicos")
add_table(["Objetivo", "Descrição"], [
    ["Privacidade", "Reduzir a observabilidade do ISP e de terceiros sobre tráfego e DNS."],
    ["Disponibilidade", "Preservar a conectividade por meio de watchdogs e validações."],
    ["Integridade", "Detectar alteração de arquivos críticos e comportamento anômalo."],
    ["Confiabilidade", "Gerar documentação reprodutível e verificável."],
], widths=[4, 12])

# =================== 2. CONTEXTO ===================
h2("2. Contexto e justificativa")
para("O contexto do projeto é a convergência entre streaming doméstico, custo mensal elevado e um "
     "mercado clandestino de TV box que opera em larga escala. Em vez de tratar a privacidade como "
     "abstração, o OpenBox parte do caso brasileiro: consumidores procuram alternativas mais baratas, "
     "marketplaces distribuem dispositivos sem garantia técnica e o ecossistema de malware explora "
     "exatamente essa cadeia de confiança frágil.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
para("O projeto também responde a uma lacuna didática. Em muitos trabalhos acadêmicos, “segurança” "
     "aparece apenas como firewall ou antivírus. Aqui, a segurança é tratada de forma sistêmica: "
     "hardware, sistema operacional, rede, monitoramento, atualização, integridade e resposta a falhas. "
     "Isso aproxima a sala de aula do que efetivamente acontece em ambientes residenciais e em pequenas "
     "infraestruturas de borda.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
callout("Panorama funcional.", "O fluxo principal é cliente → OpenBox → Internet, com o ADV7Sec "
        "observando o host e contribuindo para a telemetria local. A borda concentra as decisões de "
        "privacidade; o host concentra a detecção.", "note")

# =================== 3. LINHA DO TEMPO ===================
h2("3. Linha do tempo do projeto")
para("O OpenBox evoluiu a partir de uma base anterior de TVBox OSS e foi refinado em múltiplas etapas "
     "até atingir o formato atual. As versões iniciais consolidaram a ideia de uma caixa de borda com "
     "privacidade auditável; as versões posteriores corrigiram a documentação técnica, trocaram o alvo "
     "de hardware e incorporaram verificação estrutural do ambiente.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
para("A transição de Raspberry Pi 4 para RK3229 não representou retrocesso. Pelo contrário: o objetivo "
     "era provar que um dispositivo de baixo custo e amplamente reaproveitável pode sustentar uma pilha "
     "de segurança honesta, desde que o escopo seja bem definido. A história do projeto também registra "
     "a integração do vault Obsidian e a geração automática de documentos derivados, o que reforça "
     "rastreabilidade e manutenção.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
figure("fig2_timeline.png", "Figura 2 — Marcos de evolução, de TVBox OSS à entrega acadêmica orientada à rubrica.")
add_table(["Marco", "Descrição"], [
    ["TVBox OSS v2", "Primeira versão técnica, com inconsistências corrigidas posteriormente."],
    ["OpenBox v0.1.0", "Skeleton com WireGuard, nftables, Pi-hole, dnscrypt-proxy e watchdogs."],
    ["OpenBox v0.2.0", "Retarget para RK3229 / Armbian rk322x-box e ajuste de mídia para Jellyfin."],
    ["ADV7Box", "Consolidação em HTML+PDF com documentação expandida."],
    ["AVAL01-IC", "Reformulação acadêmica orientada à rubrica de Introdução à Computação."],
], widths=[4, 12])

doc.add_page_break()

# =================== 4. ESTUDO DE CASO ===================
h2("4. Estudo de caso: o MXQ Pro 5G 4K e a ameaça de fábrica")
para("O estudo de caso central do projeto é o MXQ Pro 5G 4K, um dos nomes genéricos mais comuns entre "
     "as TV boxes de baixo custo. A diversidade de SoCs, memórias e módulos Wi-Fi sob a mesma marca "
     "branca já é um problema: o comprador nem sempre sabe o que recebeu, o firmware pode variar entre "
     "lotes e a publicidade frequentemente infla capacidades que o hardware não entrega. O rótulo “5G” "
     "costuma referir-se apenas a uma promessa de Wi-Fi de 5 GHz, muitas vezes inexistente na placa "
     "real.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
para("A importância acadêmica desse caso está na sobreposição entre problemas de computação básica e "
     "problemas de segurança. A partir do dispositivo, é possível discutir arquitetura ARM, bootchain, "
     "módulos de comunicação, desempenho, energia, armazenamento em eMMC e, sobretudo, o risco de "
     "firmware pré-infectado. O projeto formaliza esse risco como parte da cadeia de suprimento: o "
     "problema existe antes mesmo da primeira conexão à rede.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
h3("Elementos críticos do caso")
add_table(["Elemento", "Observação"], [
    ["SoC", "Rockchip RK3229 em placa R29_5G_LP3."],
    ["Memória", "1 GB LPDDR3, o que exige tuning conservador de buffer e watchdogs."],
    ["Ethernet", "100 Mbps, suficiente para o perfil de uso e coerente com o custo."],
    ["Firmware", "Potencialmente comprometido de fábrica em determinados lotes."],
    ["Risco", "Backdoor, OTA no primeiro boot, ADB aberto, alterações não autorizadas."],
], widths=[4, 12])
callout("BadBox 2.0 / Vo1d.", "Campanhas reais infectaram milhões de dispositivos rk322x — o Brasil "
        "figura entre os mais afetados pelo Vo1d. A mitigação do OpenBox é remover o Android de fábrica "
        "por wipe via maskrom e gravar Armbian auditável antes de qualquer tráfego de rede. Detalhes em "
        "docs/security/RK3229_THREAT_RESEARCH.md.", "warn")

# =================== 5. ARQUITETURA ===================
h2("5. Arquitetura OpenBox")
para("A arquitetura de rede do OpenBox prioriza uma fronteira clara entre o que é confiável e o que não "
     "é. A unidade de borda recebe o enlace do provedor, aplica a política de saída, força a resolução "
     "DNS por um encadeamento local e encaminha o tráfego de forma encapsulada. Em paralelo, a camada de "
     "auditoria observa o comportamento do host e coleta sinais de integridade, tráfego e eventos de "
     "sistema.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
para("A topologia é compatível com uso doméstico e com pequenos laboratórios: um roteador de borda, um "
     "switch L2 com VLAN opcional, uma workstation Linux para análise e um segmento separado para "
     "dispositivos menos confiáveis, como TV boxes e IoT. O design evita complexidade desnecessária e, "
     "ao mesmo tempo, permite crescimento incremental.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
figure("fig3_flow.png", "Figura 3 — Fluxo lógico: a saída só é permitida pelo túnel; se a VPN cai, o nftables bloqueia o vazamento.")
add_table(["Etapa", "Função"], [
    ["1", "O cliente solicita resolução ou acesso."],
    ["2", "O OpenBox intercepta o DNS e aplica a política."],
    ["3", "O dnscrypt-proxy e o Unbound tratam a consulta."],
    ["4", "O WireGuard encapsula o tráfego válido."],
    ["5", "O nftables bloqueia o vazamento se a VPN cair."],
    ["6", "O ADV7Sec coleta a telemetria e alerta."],
], widths=[2.5, 13.5])

doc.add_page_break()

# =================== 6. DEMANDAS ===================
h2("6. Demandas de hardware, software e rede")
para("A demanda de hardware foi calibrada para permanecer abaixo de um orçamento doméstico razoável. O "
     "edge router pode ser um RK3229 reaproveitado ou um Raspberry Pi 4, mas a documentação final "
     "favorece o RK3229 por custo e aderência ao cenário observado. A workstation de monitoramento pode "
     "ser reaproveitada e não exige aquisição adicional, o que reforça o caráter sustentável do projeto.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
h3("Hardware recomendado")
add_table(["Componente", "Especificação alvo", "Custo aproximado"], [
    ["Edge router", "RK3229 / 1 GB RAM / eMMC 8 GB / Ethernet 100 Mbps", "R$ 180–350"],
    ["Workstation", "Linux ≥ 6.8, 4 GB RAM ou mais", "reaproveitado"],
    ["Switch VLAN", "L2 gerenciável para segmentação", "R$ 140–200"],
    ["Cabeamento", "Cat5e/6 + USB-Ethernet quando necessário", "R$ 60–120"],
], widths=[3.5, 8.5, 4])
h3("Software e serviços")
add_table(["Camada", "Ferramentas"], [
    ["Rede", "WireGuard, nftables, dnscrypt-proxy, Unbound, Pi-hole, Tor"],
    ["Monitoramento", "Netdata, Monit, Lynis, RKHunter, Uptime Kuma"],
    ["Segurança", "Fail2ban, AIDE, Suricata, Zeek, auditd, Falco"],
    ["Mídia", "Jellyfin, em substituição a opções sem suporte para armhf"],
], widths=[4, 12])
callout("Definição de infraestrutura de rede.", "A topologia mínima é provedor → edge router (OpenBox) "
        "→ switch L2/VLAN → segmentos (confiável e IoT/TV box), com a workstation de análise (ADV7Sec) "
        "em segmento confiável. A VLAN separa o que não é confiável; a VPN é a única rota de saída; o DNS "
        "é resolvido localmente e cifrado.", "note")

# =================== 7. MODELO DE AMEAÇA ===================
h2("7. Modelo de ameaça e matriz de risco")
para("O modelo de ameaça distingue o que o OpenBox mitiga do que ele não pretende resolver. ISP, "
     "ataques na LAN, telemetria proprietária e compromissos de integridade são tratados; adversários "
     "estatais com correlação global, acesso físico e compromisso do provedor de VPN permanecem fora do "
     "escopo principal. Essa delimitação é importante para evitar promessas irreais de anonimato "
     "absoluto.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
para("A matriz de risco não é apenas classificatória; ela orienta a priorização de controles. Falhas de "
     "hardware, software e rede foram mantidas em categorias distintas para preservar a lógica da "
     "rubrica e permitir mitigação específica, em vez de respostas genéricas como “instalar antivírus” "
     "ou “tomar cuidado”.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_table(["Falha", "Categoria", "Probabilidade", "Impacto", "Nota"], [
    ["Queima do gateway / falta de energia", "Hardware", "Média", "Alta", "Alto"],
    ["Erro de configuração do firewall", "Software", "Média", "Alta", "Alto"],
    ["Ataque DDoS ou interceptação de tráfego", "Rede", "Baixa", "Muito alta", "Alto"],
    ["Vazamento de DNS", "Rede", "Baixa", "Alta", "Médio/Alto"],
], widths=[5.5, 2.6, 3, 2.4, 2.5])
figure("fig4_riskmatrix.png", "Figura 4 — Matriz probabilidade × impacto. HW = hardware; SW = software; DDoS e DNS = rede.", width_in=4.6)
callout("Critério de decisão.", "A lógica do projeto é simples: se a rede cai, nada sai; se o DNS "
        "escapa, o alerta deve aparecer; se o arquivo crítico muda, a integridade é violada e o evento "
        "precisa ser registrado imediatamente.", "ok")

doc.add_page_break()

# =================== 8. HARDENING ===================
h2("8. Plano de hardening e mitigação")
para("O plano de hardening consolida a arquitetura em controles concretos. No plano de rede, o "
     "WireGuard é a única via de saída permitida; no plano de tráfego, o nftables atua como base de "
     "política e mantém o kill switch de forma atômica (via fwmark 51820). O uso de dnscrypt-proxy e "
     "Unbound reduz a superfície de observação e dificulta a manipulação de DNS por terceiros.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
para("No plano de sistema, o OpenBox adota verificações de integridade, endurecimento de SSH, "
     "segregação de serviços e atualização controlada. O plano reconhece limitações reais: o Tor não é "
     "usado para streaming HD/4K e o objetivo não é o anonimato total. Em vez disso, o Tor é reservado a "
     "metadados e consultas leves, em que o custo de latência é aceitável.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
h3("Mitigações principais")
add_table(["Falha", "Mitigação"], [
    ["Hardware", "UPS, backup de configuração, retenção de imagem de recuperação e equipamento reserva."],
    ["Software", "Git, rollback, validação de configuração e testes automatizados antes e depois da instalação."],
    ["Rede", "Firewall, VLAN, VPN obrigatória, DNS criptografado e rate limiting."],
    ["Integridade", "AIDE, RKHunter, logs e alertas por ntfy."],
], widths=[3.5, 12.5])
callout("Detalhe operacional.", "O planejamento evita composições frágeis: não se mistura UFW com "
        "nftables, não se assume que o Tor sustenta streaming de alta taxa e não se aceita um kill "
        "switch sem allowlist para o endpoint da VPN.", "warn")

# =================== 9. ADV7SEC ===================
h2("9. ADV7Sec: detecção host-side e telemetria")
para("O ADV7Sec complementa o OpenBox ao observar a workstation Linux, o que é útil quando a ameaça já "
     "passou da borda. Ele agrega eventos de auditd, Falco, Suricata e Zeek, normaliza a saída, gera "
     "alertas locais e fornece uma visão coerente de comportamento anômalo. Trata-se de um plano de "
     "controle de detecção, não de bloqueio ativo, o que simplifica a análise didática.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
para("A implementação privilegia um modo preview-first: o sistema mostra o que seria feito antes de "
     "executar ações de impacto. Esse desenho favorece estudo, depuração e aprendizado, sem transformar "
     "a plataforma em uma caixa-preta que “faz coisas” sem explicação.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
h3("Sensores e saída")
add_table(["Sensor", "O que observa", "Saída"], [
    ["auditd", "Syscalls e alterações de arquivos sensíveis", "Log estruturado"],
    ["Falco", "Comportamento de processos e rede", "Alerta e JSON"],
    ["Suricata", "Assinaturas de rede", "EVE JSON"],
    ["Zeek", "Sessões e metadados de conexão", "conn.log"],
    ["AIDE", "Integridade de arquivos", "Baseline / diff"],
], widths=[3, 8.5, 4.5])

doc.add_page_break()

# =================== 10. RUNBOOK ===================
h2("10. Runbook, validação e reprodutibilidade")
para("A qualidade operacional do OpenBox depende de uma rotina clara de verificação. O runbook define "
     "comandos rápidos para checar serviços, conexões, vazamento de DNS, circuito Tor e integridade "
     "geral. Em sistemas de segurança, a documentação operacional é parte da solução, porque reduz o "
     "tempo entre um sintoma e a correção.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
h3("Healthchecks resumidos")
add_table(["Checagem", "Comando de referência"], [
    ["Estado geral", "systemctl status wg-quick@wg0 tor dnscrypt-proxy pihole-FTL nftables netdata"],
    ["IP de saída", "curl -4 https://ifconfig.io"],
    ["Circuito Tor", "curl --socks5-hostname 127.0.0.1:9050 https://check.torproject.org/api/ip"],
    ["DNS leak", "openbox-dnsleak-check.sh"],
    ["Hardening index", "lynis audit system --quick --no-colors"],
], widths=[3.5, 12.5])
h3("Falhas comuns e resposta prevista")
add_table(["Falha", "Resposta prevista"], [
    ["Handshake travado", "Reiniciar wg-quick@wg0 e verificar latest-handshakes."],
    ["DNSCrypt sem resposta", "Verificar o bind em 127.0.0.1:5053 e o upstream do Pi-hole."],
    ["Painel Netdata não abre", "Validar o Caddy, o reverse proxy e os logs locais."],
    ["Tor degradado", "Reiniciar o serviço e inspecionar o circuito com nyx."],
], widths=[4, 12])

# =================== 11. EVIDÊNCIAS ===================
h2("11. Evidências de desenvolvimento e rastreabilidade")
para("A entrega acadêmica é acompanhada por documentação gerada e versionada. O repositório inclui "
     "README, changelog, documentação de ameaça, guia de setup de hardware, runbook, referências e o "
     "vault do Obsidian com mapas de conteúdo gerados automaticamente. Isso permite auditoria posterior "
     "e reduz a dependência de texto manual desconectado do projeto real.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
para("Em vez de tratar documentos como anexos decorativos, o OpenBox os integra como parte da "
     "engenharia. Os artefatos textuais descrevem as fases de instalação, os arquivos de configuração e "
     "o caminho de verificação; as notas do vault mantêm a trilha de execução; e os scripts de build e "
     "validação preservam a reprodutibilidade.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
h3("Artefatos mais relevantes")
add_table(["Artefato", "Função"], [
    ["README.md", "Visão geral do projeto e diferenças em relação ao TVBox OSS."],
    ["CHANGELOG.md", "Histórico de mudanças entre as versões 0.1.0 e 0.2.0."],
    ["THREAT_MODEL.md", "Adversários, fronteiras de confiança e modos de falha."],
    ["RK3229_THREAT_RESEARCH.md", "Ameaças de cadeia de suprimento e mitigação."],
    ["RUNBOOK.md", "Operações diárias e troubleshooting."],
    ["OBSIDIAN_VAULT.md", "Integração do repositório como vault do Obsidian."],
], widths=[5, 11])

doc.add_page_break()

# =================== 12. CONCLUSÃO ===================
h2("12. Conclusão e direcionamento de carreira")
para("O OpenBox v0.2.0 mostra que uma solução defensiva bem definida pode ser construída com hardware "
     "de baixo custo, software livre e disciplina metodológica. O valor do projeto está menos em "
     "prometer uma abstração de “segurança máxima” e mais em transformar práticas concretas — segmentar, "
     "observar, validar, bloquear e registrar — em uma rotina acessível e replicável.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
para("No plano pessoal e profissional, o projeto aponta para a atuação em infraestrutura, redes e "
     "segurança da informação, com forte viés Blue Team. A combinação entre hardening, observabilidade, "
     "resposta a incidentes e documentação técnica é compatível com carreiras em SOC, administração de "
     "redes, segurança de sistemas e engenharia de plataformas. A base construída aqui pode evoluir para "
     "maior automação, análise de alertas e integração com ferramentas de inteligência local em "
     "semestres futuros.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
callout("Fechamento.", "Em termos acadêmicos, o projeto atende ao pedido de estudar um cenário real, "
        "justificar a tecnologia e demonstrar uma solução defensiva. Em termos práticos, ele entrega uma "
        "borda residencial mais segura, compreensível e auditável.", "ok")

# =================== 13. REFERÊNCIAS ===================
h2("13. Referências técnicas essenciais")
add_table(["Categoria", "Exemplos de referência"], [
    ["VPN", "WireGuard whitepaper, guias de hardening, documentação de tuning."],
    ["Firewall", "nftables, kill switch com fwmark, allowlist de endpoint."],
    ["DNS", "Pi-hole, dnscrypt-proxy, DNSSEC e DoH/DoQ."],
    ["Monitoramento", "Netdata, Monit, Lynis, RKHunter."],
    ["IDS/IPS", "Suricata, Zeek, auditd, Falco."],
    ["Mídia", "Jellyfin para armhf; Tor apenas para metadados."],
], widths=[4, 12])
para("A bibliografia técnica do repositório já registra referências verificadas para as escolhas de "
     "rede, hardening, monitoramento e mídia. A integração desses materiais ao projeto garante que as "
     "decisões não são arbitrárias, mas derivadas de documentação e validação operacional.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

doc.add_page_break()

# =================== 14. CHECKLIST ===================
h2("14. Checklist final de conformidade")
add_table(["Requisito da AVFinal", "Atendido em", "Status"], [
    ["Definição de infraestrutura de redes", "Seções 5 e 6", "✓"],
    ["Exatamente 3 falhas (hardware, software, rede)", "Seção 7", "✓"],
    ["Plano de mitigação", "Seção 8", "✓"],
    ["Evidências de desenvolvimento", "Seção 11", "✓"],
    ["Parágrafo de área de atuação", "Seção 12", "✓"],
    ["Critérios de avaliação / rubrica", "Estrutura incorporada ao raciocínio do documento", "✓"],
], widths=[8, 5.5, 2.5])
para("Este anexo resume a correspondência entre a rubrica e o documento final. O trabalho foi "
     "organizado para que cada requisito seja encontrável sem adivinhação e com um fluxo de leitura "
     "coerente.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
figure("fig5_closing.png", "Figura 5 — Defesa em profundidade: rede, DNS, VPN, monitoramento e integridade.")
para("Documento final gerado em ambiente local com base nos arquivos do projeto OpenBox0.1v, incluindo "
     "README, changelog, threat model, runbook, pesquisa de hardware e vault Obsidian.",
     size=9, color=MUTED, after=14)

# =================== ANEXO A — TEARDOWN + FLASH/SSH ===================
doc.add_page_break()
h2("Anexo A — Teardown de hardware e fase de flash/SSH")
para("Este anexo documenta a inspeção física da unidade real, identificando as estruturas da placa e "
     "registrando a fase de flash e o primeiro acesso por SSH. A placa foi identificada pela serigrafia "
     "como R3290_V8.1 (data 2020.06.15, marca de teste TY 2324), da família rk322x (RK3228A/B ou "
     "RK3229), coberta pela imagem Armbian rk322x-box.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
h3("A.1 Identificação da placa")
add_table(["Item", "Valor observado"], [
    ["Serigrafia da placa", "R3290_V8.1 · 2020.06.15 · marca de teste TY 2324"],
    ["Família / SoC", "rk322x — Rockchip RK3229 (ARM Cortex-A7 quad), sob dissipador"],
    ["Armazenamento", "eMMC Samsung KLM8G1GETF — 8 GB (BGA-153)"],
    ["Memória", "4× Hynix H5TQ2G43BFR (DDR3, 2 por face) ≈ 1 GB"],
    ["Imagem alvo", "Armbian community rk322x-box (Debian Bookworm, armhf)"],
], widths=[4, 12])
h3("A.2 Estruturas identificadas")
add_table(["#", "Estrutura", "Marcação / parte", "Função"], [
    ["1", "eMMC", "Samsung KLM8G1GETF (SEC 219 B041)", "Armazenamento principal — 8 GB eMMC 5.x"],
    ["2", "R-série", "resistores em série", "Tap/isolação no barramento da eMMC"],
    ["3", "Trilha CLK", "breakout junto à eMMC", "Short para entrar em maskrom (método B)"],
    ["4", "CI companion", "QFN-44 (SVS6Z56F / TAC2121)", "PHY/controlador de interface (provável)"],
    ["5", "PPT PM44-11BP", "SOIC junto a RJ45/DC", "MOSFET/driver de potência (provável)"],
    ["6", "Cap. eletrolítico", "CK 100 µF / 10 V", "Filtro da entrada de energia"],
    ["7", "LDO", "AMS1117 (1117C, SOT-223)", "Regulador linear 3,3 V"],
    ["8", "Conversores buck", "2× indutor 3R3 (3,3 µH)", "Trilhos de núcleo/DDR"],
    ["9", "SoC", "Rockchip RK3229 (sob dissipador)", "CPU ARM Cortex-A7 quad — rk322x"],
    ["10", "Header serial", "R / T / G (RX, TX, GND)", "Console UART de depuração"],
    ["11–12", "DRAM", "Hynix H5TQ2G43BFR (DDR3)", "4 chips (2 por face) ≈ 1 GB"],
    ["13", "Rede de resistores", "face inferior", "Terminação/strap do barramento DDR"],
    ["14", "LDO auxiliar", "1117 (face inferior)", "Regulador secundário"],
    ["15", "Footprint BGA vazio", "não populado", "Posição alternativa de NAND/eMMC"],
    ["16", "Cristal/oscilador", "—", "Referência de clock (24/25 MHz)"],
    ["17", "USB-A", "—", "Porta host USB 2.0"],
    ["18–19", "AV 3,5 mm", "—", "Saídas CVBS/áudio"],
    ["20", "HDMI", "—", "Saída de vídeo digital"],
    ["21", "RJ45", "—", "Ethernet 10/100"],
    ["22", "DC jack", "—", "Entrada de alimentação 5 V"],
    ["—", "Micro-USB OTG", "topo", "Porta usada para o flash via maskrom"],
    ["—", "IR · microSD/TF", "frontal · lateral", "Controle remoto · armazenamento removível"],
], widths=[1.3, 3.2, 4.5, 7.0])
figure("hw/plate_top.png", "Figura A.1 — Face superior: eMMC, SoC (sob dissipador), regulação e header serial.", width_in=3.0)
figure("hw/plate_bottom.png", "Figura A.2 — Face inferior: DRAM Hynix, footprint BGA vazio e LDO. MAC redigido por privacidade.", width_in=6.0)
figure("hw/plate_ports.png", "Figura A.3 — Conectividade: USB-A, AV 3,5 mm, HDMI, RJ45 e DC 5 V.", width_in=6.0)
h3("A.3 Fase de flash e SSH")
para("O procedimento de gravação foi conduzido a partir do host de desenvolvimento (Windows + WSL Kali), "
     "anexando o dispositivo USB com usbipd-win e gravando a imagem Armbian com rkdeveloptool em modo "
     "maskrom (PID 2207:320b). Após o boot, o primeiro acesso ocorreu por SSH com autenticação por chave, "
     "seguido do gate de fingerprint da Fase 0 (SoC, arquitetura, módulo WireGuard e /dev/watchdog).",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
callout("Reconstrução fiel.", "Os terminais reproduzem o procedimento documentado com os valores reais do "
        "projeto (placa R3290_V8.1 · Armbian rk322x-box · host 192.168.0.103 · Phase 0: PASS 4/4), a partir "
        "do runbook e das notas de execução.", "note")
figure("hw/term_flash.png", "Figura A.4 — Fase de flash: maskrom → rkdeveloptool wl → rd.", width_in=6.3)
figure("hw/term_ssh.png", "Figura A.5 — Fase de SSH: primeiro acesso ao Armbian e gate da Fase 0.", width_in=6.3)

# signature row
sigt = doc.add_table(rows=1, cols=2)
for i, (a, b) in enumerate([("Alcides Olivo Pollazzon Soterio", "Autor"),
                            ("UniCEUB · Introdução à Computação", "Brasília – DF · 2026")]):
    c = sigt.rows[0].cells[i]
    p1 = c.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(a); r1.font.size = Pt(9.5); r1.bold = True
    p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(b); r2.font.size = Pt(9); r2.font.color.rgb = MUTED

# footer
footer = doc.sections[0].footer
fp = footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("OpenBox AVFinal · UniCEUB · Introdução à Computação · 2026")
fr.font.size = Pt(8); fr.font.color.rgb = MUTED

out = HERE / "OpenBox_AVFinal.docx"
doc.save(str(out))
print("DOCX saved:", out, out.stat().st_size, "bytes")
